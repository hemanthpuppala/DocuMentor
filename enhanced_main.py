# main.py - DocuMentor Backend (Production-Grade Version)

# === Core Imports ===
import os
import re
import requests
import asyncio
import tempfile
import traceback
from typing import List, Dict, Any, Literal, Optional
from pydantic import BaseModel, ValidationError
from typing_extensions import TypedDict
from dotenv import load_dotenv
from datetime import datetime
from docx import Document
import subprocess
import platform


# === LangChain & LangGraph ===
from langchain_groq import ChatGroq
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import StateGraph, START, END


# === Document Processing ===
import fitz  # PyMuPDF
import arxiv


# === Web Framework ===
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import uuid
import shutil
# === Web Framework ===
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request, APIRouter  # ADD APIRouter here
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi import APIRouter


# === Environment & Config ===
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
BRAVE_API_KEY = os.getenv("BRAVE_API_KEY")


# === Constants ===
PRIMARY_MODEL_NAME = "gemini-1.5-flash"
FALLBACK_MODEL_NAME = "gemma2-9b-it" #"deepseek-r1-distill-llama-70b" 
MAX_DOCUMENT_SIZE = 10 * 1024 * 1024  # 10MB
MAX_TEXT_LENGTH = 50000  # characters
ALLOWED_FILE_TYPES = {'.pdf', '.txt', '.doc', '.docx'}


# === LLM Initialization ===
primary_llm = None
fast_llm = None


def initialize_llms():
    global primary_llm, fast_llm
    try:
        if GOOGLE_API_KEY:
            primary_llm = ChatGoogleGenerativeAI(
                model=PRIMARY_MODEL_NAME, 
                temperature=0.3, 
                google_api_key=GOOGLE_API_KEY, 
                convert_system_message_to_human=True
            )
            fast_llm = ChatGoogleGenerativeAI(
                model=PRIMARY_MODEL_NAME, 
                temperature=0.1, 
                google_api_key=GOOGLE_API_KEY, 
                convert_system_message_to_human=True
            )
            print(f"✅ Using Google Gemini: {PRIMARY_MODEL_NAME}")
        elif GROQ_API_KEY:
            primary_llm = ChatGroq(
                model=FALLBACK_MODEL_NAME, 
                groq_api_key=GROQ_API_KEY, 
                temperature=0.3
            )
            fast_llm = ChatGroq(
                model=FALLBACK_MODEL_NAME, 
                groq_api_key=GROQ_API_KEY, 
                temperature=0.1
            )
            print(f"⚠️ Google API key not found. Using Groq fallback: {FALLBACK_MODEL_NAME}")
        else:
            raise ValueError("No LLM API key found. Please set GOOGLE_API_KEY or GROQ_API_KEY.")
        return True
    except Exception as e:
        print(f"❌ LLM initialization failed: {e}")
        return False


# === State Management ===
class ResearchState(TypedDict):
    original_question: str
    generated_question: str
    formatted_question: str
    session_id: str
    uploaded_documents_raw: List[str]
    internet_results_raw: List[str]
    arxiv_results_raw: List[str]
    source_summaries: List[str]
    research_summary: str
    user_question: str
    specialist_response: str
    errors: List[str]
    processing_stage: str
    chat_history: List[Dict[str, str]]  # NEW: Store chat conversation history
    session_metadata: Dict[str, Any]    # NEW: Store session info (title, created_at, etc.)
    latest_quiz_html: str            # NEW – last quiz sent to user
    latest_quiz_time: str            # NEW – ISO timestamp
    quiz_history: List[Dict[str, Any]]   # NEW – list of past attempts
    follow_up_question: str  # NEW – last follow-up question asked by user


session_store: Dict[str, ResearchState] = {}


# === Utility Functions ===
def safe_extract_text_from_pdf(file_path: str) -> str:
    """Safely extract text from PDF with error handling"""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page_num in range(len(doc)):
            try:
                page = doc[page_num]
                page_text = page.get_text()
                text += page_text + "\n"
            except Exception as e:
                print(f"⚠️ Error reading page {page_num}: {e}")
                continue
        doc.close()
        return text.strip()
    except Exception as e:
        print(f"❌ PDF extraction failed: {e}")
        return ""


def safe_extract_text_from_file(file_path: str) -> str:
    """Safely extract text from various file types"""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        
        if file_ext == '.pdf':
            return safe_extract_text_from_pdf(file_path)
            
        elif file_ext == '.docx':
            # Handle DOCX files using python-docx
            try:
                doc = Document(file_path)
                full_text = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text)
                
                text = '\n'.join(full_text)
                
                # Apply length limit
                if len(text) > MAX_TEXT_LENGTH:
                    text = text[:MAX_TEXT_LENGTH] + "\n[Content truncated...]"
                
                print(f"✅ Successfully extracted text from DOCX: {filename}")
                return text.strip()
                
            except Exception as e:
                print(f"❌ DOCX extraction failed for {filename}: {e}")
                return f"Error extracting text from DOCX file: {filename}"
                
        elif file_ext == '.doc':
            # Handle legacy DOC files - multiple approaches
            try:
                # Method 1: Try using python-docx (sometimes works with DOC)
                try:
                    doc = Document(file_path)
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text)
                    
                    if full_text:
                        text = '\n'.join(full_text)
                        if len(text) > MAX_TEXT_LENGTH:
                            text = text[:MAX_TEXT_LENGTH] + "\n[Content truncated...]"
                        print(f"✅ Successfully extracted text from DOC: {filename}")
                        return text.strip()
                except:
                    pass
                
                # Method 2: Try using textract (if available)
                try:
                    import textract
                    text = textract.process(file_path).decode('utf-8')
                    if len(text) > MAX_TEXT_LENGTH:
                        text = text[:MAX_TEXT_LENGTH] + "\n[Content truncated...]"
                    print(f"✅ Successfully extracted text from DOC using textract: {filename}")
                    return text.strip()
                except ImportError:
                    pass
                except Exception as e:
                    print(f"⚠️ Textract extraction failed: {e}")
                
                # Method 3: System-specific conversion
                if platform.system() == "Windows":
                    try:
                        # Try using COM objects on Windows
                        import win32com.client
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        doc = word.Documents.Open(file_path)
                        text = doc.Content.Text
                        doc.Close()
                        word.Quit()
                        
                        if len(text) > MAX_TEXT_LENGTH:
                            text = text[:MAX_TEXT_LENGTH] + "\n[Content truncated...]"
                        print(f"✅ Successfully extracted text from DOC using COM: {filename}")
                        return text.strip()
                    except Exception as e:
                        print(f"⚠️ COM extraction failed: {e}")
                
                # Fallback message
                return f"Legacy DOC format detected in {filename}. For best results, please convert to DOCX format. Some text may be available but extraction is limited."
                
            except Exception as e:
                print(f"❌ DOC extraction failed for {filename}: {e}")
                return f"Error processing DOC file: {filename}. Please convert to DOCX for better compatibility."
                
        elif file_ext == '.txt':
            # Handle text files with multiple encoding attempts
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        if len(content) > MAX_TEXT_LENGTH:
                            content = content[:MAX_TEXT_LENGTH] + "\n[Content truncated...]"
                        print(f"✅ Successfully extracted text from TXT: {filename} (encoding: {encoding})")
                        return content
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    print(f"⚠️ Error with encoding {encoding}: {e}")
                    continue
            
            # If all encodings fail, try with error handling
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    if len(content) > MAX_TEXT_LENGTH:
                        content = content[:MAX_TEXT_LENGTH] + "\n[Content truncated...]"
                    print(f"⚠️ Text extracted with encoding issues: {filename}")
                    return content
            except Exception as e:
                print(f"❌ Final text extraction attempt failed: {e}")
                return f"Error reading text file: {filename}"
        else:
            error_msg = f"Unsupported file type: {file_ext}. Supported: {', '.join(ALLOWED_FILE_TYPES)}"
            print(f"❌ {error_msg}")
            return ""
            
    except Exception as e:
        error_msg = f"File processing failed for {filename if 'filename' in locals() else file_path}: {str(e)}"
        print(f"❌ {error_msg}")
        return ""


def safe_internet_search_detailed(query: str, max_results: int = 5) -> List[str]:
    """Return richer snippets instead of only title/URL."""
    if not BRAVE_API_KEY:
        print("⚠️  BRAVE_API_KEY not set")
        return []

    try:
        headers = {
            "Accept": "application/json",
            "Accept-Encoding": "gzip",
            "X-Subscription-Token": BRAVE_API_KEY
        }
        params = {"q": query, "count": max_results, "offset": 0,
                  "mkt": "en-US", "safesearch": "moderate"}
        r = requests.get("https://api.search.brave.com/res/v1/web/search",
                         headers=headers, params=params, timeout=15)
        if r.status_code != 200:
            print(f"⚠️ Brave API status {r.status_code}")
            return []

        data = r.json()
        results = []
        for item in data.get("web", {}).get("results", []):
            parts = []
            if item.get("title"):
                parts.append(f"<h3>{item['title']}</h3>")
            if item.get("description"):
                parts.append(f"<p><strong>Overview:</strong> {item['description']}</p>")
            if item.get("snippet"):
                parts.append(f"<p>{item['snippet']}</p>")
            if parts:
                results.append("\n".join(parts))
        print(f"🌐 Internet search found {len(results)} rich results")
        return results
    except Exception as e:
        print(f"❌ Brave search error: {e}")
        return []


async def extract_detailed_arxiv_content(papers) -> List[str]:
    detailed = []
    for paper in papers:
        try:
            blocks = [
                f"<h2>{paper.title}</h2>",
                f"<p><strong>Authors:</strong> {', '.join(a.name for a in paper.authors)}</p>",
                f"<p><strong>Published:</strong> {paper.published:%Y-%m-%d}</p>",
                f"<p><strong>Category:</strong> {paper.primary_category}</p>",
                "<h3>Abstract</h3>",
                f"<p>{paper.summary}</p>",
            ]
            blocks.extend(analyze_abstract_for_sections(paper.summary))
            detailed.append("\n".join(blocks))
        except Exception as e:
            print(f"❌ Extract error on '{paper.title}': {e}")
    return detailed



def analyze_abstract_for_sections(abstract: str) -> List[str]:
    sections, sentences = [], abstract.split('. ')
    if len(sentences) >= 3:
        sections.append("<h3>Research Focus</h3>")
        sections.append(f"<p>{'. '.join(sentences[:2])}.</p>")
        sections.append("<h3>Key Contributions</h3>")
        sections.append(f"<p>{'. '.join(sentences[2:4])}.</p>")
    if len(sentences) > 4:
        sections.append("<h3>Methodology Overview</h3>")
        sections.append(f"<p>{'. '.join(sentences[4:])}.</p>")
    return sections


async def safe_arxiv_search_detailed(query: str, max_results: int = 3) -> List[str]:
    try:
        client = arxiv.Client()
        search = arxiv.Search(query=query, max_results=max_results,
                              sort_by=arxiv.SortCriterion.Relevance)
        papers = list(client.results(search))
        if not papers:
            print(f"📚 No papers for '{query}'")
            return []
        content = await extract_detailed_arxiv_content(papers)
        print(f"📚 ArXiv returned {len(content)} detailed papers for '{query}'")
        return content
    except Exception as e:
        print(f"❌ ArXiv search error '{query}': {e}")
        return []

async def safe_llm_invoke(llm, prompt_data: dict, max_retries: int = 3) -> str:
    """Safely invoke LLM with retries and error handling"""
    for attempt in range(max_retries):
        try:
            response = await llm.ainvoke(prompt_data)
            if hasattr(response, 'content'):
                return response.content.strip()
            else:
                return str(response).strip()
        except Exception as e:
            print(f"⚠️ LLM invoke attempt {attempt + 1} failed: {e}")
            if attempt == max_retries - 1:
                print(f"❌ All LLM attempts failed")
                return f"Error: Unable to process request - {str(e)}"
            await asyncio.sleep(1)  # Wait before retry
    return "Error: Maximum retries exceeded"


# === LangGraph Nodes (The "Agents") ===

async def document_processor_node(state: ResearchState) -> ResearchState:
    """Process uploaded documents and extract text"""
    print("📄 Step 1a: Document Processor - extracting text...")
    state["processing_stage"] = "document_processing"
    
    try:
        if state.get("uploaded_documents_raw"):
            print(f"📄 Processing {len(state['uploaded_documents_raw'])} documents")
            # Documents are already processed in the API, just validate here
            valid_docs = []
            for doc in state["uploaded_documents_raw"]:
                if doc and doc.strip():
                    valid_docs.append(doc)
            
            state["uploaded_documents_raw"] = valid_docs
            print(f"📄 Validated {len(valid_docs)} documents")
        else:
            print("📄 No documents to process")
            
    except Exception as e:
        print(f"❌ Document processing error: {e}")
        state["errors"].append(f"Document processing failed: {str(e)}")
    
    return state


async def query_generator_node(state: ResearchState) -> ResearchState:
    """Generate a research query from document content"""
    print("💡 Step 1b: Query Generator - analyzing documents...")
    state["processing_stage"] = "query_generation"
    
    try:
        if not state.get("uploaded_documents_raw"):
            print("💡 No documents found, skipping query generation")
            return state
        
        doc_content = "\n\n".join(state["uploaded_documents_raw"])[:4000]  # Limit context
        
        prompt_template = ChatPromptTemplate.from_messages([
            ("system", """You are a research assistant. Based on the user's question and the document content, 
            generate a single, specific, and clear research question. The question should be:
            1. Directly answerable from the document
            2. Suitable for further web searches
            3. Clear and specific
            
            Provide only the research question, nothing else."""),
            ("human", "User's question: '{original_question}'\n\nDocument content:\n---\n{document_content}\n\n---\n\nGenerated research question:")
        ])
        
        chain = prompt_template | fast_llm
        prompt_data = {
            "original_question": state["original_question"], 
            "document_content": doc_content
        }
        
        response_content = await safe_llm_invoke(chain, prompt_data)
        
        if response_content and not response_content.startswith("Error:"):
            state["generated_question"] = response_content
            print(f"💡 Generated Question: {state['generated_question']}")
        else:
            print(f"⚠️ Query generation failed, using original question")
            state["errors"].append("Failed to generate question from document")
            
    except Exception as e:
        print(f"❌ Query generation error: {e}")
        state["errors"].append(f"Query generation failed: {str(e)}")
    
    return state


async def formatter_node(state: ResearchState) -> ResearchState:
    """Format and enhance the research query"""
    print("🔧 Step 2: Formatter - enhancing query...")
    state["processing_stage"] = "query_formatting"
    try:
        # Check if this is a follow-up query or initial research query
        user_question = state.get("user_question", "").strip()
        
        if user_question:
            # This is a follow-up question - use the user's follow-up question
            question_to_format = user_question
            print(f"🔧 Using follow-up question: {question_to_format}")
        else:
            # This is an initial research query - use generated question or original
            question_to_format = state.get("generated_question") or state["original_question"]
            print(f"🔧 Using initial research question: {question_to_format}")

        document_context = ""
        if state.get("uploaded_documents_raw"):
            # Summarize uploaded documents for context
            docs_preview = "\n".join(state["uploaded_documents_raw"])[:1000]
            document_context = f"Available documents contain: {docs_preview}..."
        elif state.get("generated_question"):
            document_context = f"Generated from documents: {state['generated_question']}"
        
        prompt_template = ChatPromptTemplate.from_messages([
            ("system", """You are a Query Enhancement Specialist. Transform the user's question into 
            a well-structured research query suitable for web searches and arxiv search at a time. 
            Make it:
            1. Clear and specific
            2. Search-engine friendly
            
            Provide only the enhanced query, nothing else."""),
            ("human", "Original query: {question_to_format}"),
            ("human", "Document context:\n---\n{document_context}\n\n---\n\nEnhanced query:")
        ])
        
        chain = prompt_template | fast_llm
        prompt_data = {"question_to_format": question_to_format,
                       "document_context": document_context,
                       }
        
        response_content = await safe_llm_invoke(chain, prompt_data)
        
        if response_content and not response_content.startswith("Error:"):
            state["formatted_question"] = response_content
            print(f"🔧 Formatted Question: {state['formatted_question']}")
        else:
            print(f"⚠️ Query formatting failed, using original")
            state["formatted_question"] = question_to_format
            state["errors"].append("Query formatting had issues")
            
    except Exception as e:
        print(f"❌ Query formatting error: {e}")
        state["formatted_question"] = state.get("generated_question") or state["original_question"]
        state["errors"].append(f"Query formatting failed: {str(e)}")
    
    return state
async def enhanced_formatter_node(state: ResearchState) -> ResearchState:
    """Format and enhance the research query"""
    print("🔧 Step 2: Enhanced Formatter - enhancing query with full context...")
    state["processing_stage"] = "enhanced_query_formatting"
    
    try:
        # Extract document context
        document_context = ""
        if state.get("uploaded_documents_raw"):
            # Summarize uploaded documents for context
            docs_preview = "\n".join(state["uploaded_documents_raw"])[:1000]
            document_context = f"Available documents contain: {docs_preview}..."
        elif state.get("generated_question"):
            document_context = f"Generated from documents: {state['generated_question']}"
        
        # Extract and format chat history
        chat_history_context = ""
        if state.get("chat_history") and len(state["chat_history"]) > 0:
            recent_chat = state["chat_history"][-3:]  # Last 3 exchanges
            chat_entries = []
            for msg in recent_chat:
                role = msg.get("role", "user")
                content = msg.get("content", "")[:200]  # First 200 chars
                chat_entries.append(f"{role}: {content}")
            chat_history_context = "\n".join(chat_entries)
        
        # Use generated question if available, otherwise use original
        question_to_format = state.get("generated_question") or state["original_question"]
        
        prompt_template = ChatPromptTemplate.from_messages([
            ("system", """
            You are a Query Enhancement Specialist. Transform the user's question into a well-structured research query suitable for web searches and arxiv paper search at a time. 
            Make sure the formatted query is directly related to the actual user's question, chat history and the context provided.
            Also consider the recent conversation history to ensure the query is relevant and builds on previous discussions.
            If the question is a follow-up and is vague, clarify it based on the previous 🤖 Assistant message, and context to ensure it is specific to context and clear.
            Make it:
            1. Clear and specific
            2. Search-engine friendly
            3. Directly related to the user's question and context
            4. Concise and focused
            
            Provide only the enhanced query, nothing else."""),
            ("human", """Original query: {question_to_format}

            Document context: {document_context}

            Recent conversation: {chat_history_context}

            Enhanced query:
            """
             )
        ])

        chain = prompt_template | fast_llm
        prompt_data = {
            "question_to_format": question_to_format,
            "document_context": document_context,
            "chat_history_context": chat_history_context
        }
        
        response_content = await safe_llm_invoke(chain, prompt_data)
        
        if response_content and not response_content.startswith("Error:"):
            state["formatted_question"] = response_content
            print(f"🔧 Formatted Question: {state['formatted_question']}")
        else:
            print(f"⚠️ Query formatting failed, using original")
            state["formatted_question"] = question_to_format
            state["errors"].append("Query formatting had issues")
        print(f"🔧 Enhanced formatted Question: {state['formatted_question']}")
    except Exception as e:
        print(f"❌ Query formatting error: {e}")
        state["formatted_question"] = state.get("generated_question") or state["original_question"]
        state["errors"].append(f"Query formatting failed: {str(e)}")
    
    return state

async def internet_search_node(state: ResearchState) -> ResearchState:
    """Use rich Brave snippets."""
    print("🌐 Step 3 – Internet Search")
    state["processing_stage"] = "internet_search"

    try:
        primary = state.get("formatted_question", state["original_question"])
        fallback = state["original_question"]

        results = safe_internet_search_detailed(primary, 5)
        if not results and primary != fallback:
            results = safe_internet_search_detailed(fallback, 5)
        if not results:
            short = " ".join(fallback.split()[:5])
            results = safe_internet_search_detailed(short, 5)

        state["internet_results_raw"] = results
        print(f"🌐 Stored {len(results)} internet results")
    except Exception as e:
        print(f"❌ Internet search node error: {e}")
        state["internet_results_raw"] = []
        state["errors"].append(str(e))

    return state


async def arxiv_search_node(state: ResearchState) -> ResearchState:
    """Async ArXiv search without asyncio.run conflicts."""
    print("📚 Step 4 – ArXiv Search")
    state["processing_stage"] = "arxiv_search"

    try:
        primary = state.get("formatted_question", state["original_question"])
        fallback = state["original_question"]

        results = await safe_arxiv_search_detailed(primary, 3)
        if not results and primary != fallback:
            results = await safe_arxiv_search_detailed(fallback, 3)
        if not results:
            key = extract_academic_keywords(fallback)
            if key != fallback:
                results = await safe_arxiv_search_detailed(key, 3)
        if not results:
            simple = " ".join(fallback.split()[:3])
            results = await safe_arxiv_search_detailed(simple, 3)

        state["arxiv_results_raw"] = results
        if results:
            print(f"📚 Stored {len(results)} ArXiv papers")
        else:
            state["errors"].append("No relevant academic papers found")
    except Exception as e:
        print(f"❌ ArXiv node error: {e}")
        state["arxiv_results_raw"] = []
        state["errors"].append(str(e))

    return state


def extract_academic_keywords(query: str) -> str:
    """Extract academic keywords from a query"""
    # Remove common words and focus on academic terms
    stop_words = {'the', 'is', 'are', 'what', 'how', 'why', 'when', 'where', 'who', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
    
    words = query.lower().split()
    academic_words = [word for word in words if word not in stop_words and len(word) > 2]
    
    # Take up to 4 most relevant terms
    return " ".join(academic_words[:4])


async def handle_insufficient_data_node(state: ResearchState) -> ResearchState:
    """Handle cases where insufficient data was found"""
    print("🚦 Step 5a: Handling Insufficient Data...")
    state["processing_stage"] = "insufficient_data_handling"
    
    try:
        # Check what sources we have
        has_docs = bool(state.get("uploaded_documents_raw"))
        has_internet = bool(state.get("internet_results_raw"))
        has_arxiv = bool(state.get("arxiv_results_raw"))
        
        if not has_docs and not has_internet and not has_arxiv:
            error_message = "No information could be retrieved from any source for this query."
            state["source_summaries"] = [error_message]
        else:
            # We have some data, prepare what we have
            summaries = []
            
            if has_docs:
                summaries.extend(state["uploaded_documents_raw"])
                
            if has_internet:
                summaries.extend(state["internet_results_raw"])
                
            if has_arxiv:
                summaries.extend(state["arxiv_results_raw"])
            
            state["source_summaries"] = summaries
            
            error_message = "Limited search results found. Response based on available sources."
            
        state["errors"].append(error_message)
        print(f"🚦 {error_message}")
        
    except Exception as e:
        print(f"❌ Insufficient data handling error: {e}")
        state["errors"].append(f"Error handling insufficient data: {str(e)}")
        state["source_summaries"] = ["Error: Unable to process available data"]
    
    return state


async def summarize_sources_node(state: ResearchState) -> ResearchState:
    """Summarize information from all sources"""
    print("✍️ Step 5b: Summarizing Sources...")
    state["processing_stage"] = "source_summarization"
    
    try:
        all_sources = []
        
        # Collect all available sources
        if state.get("uploaded_documents_raw"):
            all_sources.extend(state["uploaded_documents_raw"])
            
        if state.get("internet_results_raw"):
            all_sources.extend(state["internet_results_raw"])
            
        if state.get("arxiv_results_raw"):
            all_sources.extend(state["arxiv_results_raw"])
        
        if not all_sources:
            state["source_summaries"] = ["No sources available for summarization"]
            return state
        
        # Limit content length
        content = "\n\n---SOURCE BREAK---\n\n".join(all_sources)
        if len(content) > 15000:  # Larger limit for detailed reports
            content = content[:15000] + "\n\n[Content truncated...]"
        
        prompt_template = ChatPromptTemplate.from_messages([
            ("system", """
You are an expert academic researcher and university-level educator. Your role is to generate detailed and structured research reports that teach the topic to an intelligent non-expert, based on provided materials.

You have access to trusted research inputs, including internet results, academic papers (from arXiv), and any uploaded user documents. Your task is to synthesize these into a single, accurate, and pedagogical report.

OBJECTIVE: Produce a full-length, structured research report that explains the topic comprehensively. The report must be accurate, easy to understand, and written in the style of a professor explaining the topic to university students or a non-specialist audience.

STRATEGY: Use chain-of-thought reasoning to unpack complex ideas step-by-step. Use multi-shot logic: compare academic perspectives, summarize multiple viewpoints. Avoid summarizing too briefly. Teach the full topic like a classroom lesson. Use real-world analogies and examples to explain abstract concepts.

Only use information from provided internet data, academic papers, uploaded documents, or generally accepted academic knowledge. Highlight any disagreements or uncertainties among sources. Do not hallucinate or fabricate any information or sources.

Use clear and explanatory language, assuming no prior expertise.

OUTPUT FORMAT: Use clean, semantic HTML and follow this strict section structure. Each section must be present and detailed:

RESTRICTIONS: Do not fabricate any sources or papers. Do not use the user's question verbatim. Do not include vague, filler text. Do not repeat sections or summarize prematurely. Do not omit any section unless data is completely unavailable (in which case state this explicitly).

Always write in full, explanatory paragraphs unless a bullet list improves clarity. Avoid technical jargon unless clearly defined first.

TONE: Academic, professional, and instructional. Prioritize clarity, factual correctness, and completeness. Teach as if speaking to a smart but non-expert audience.

You must use the provided structured data as input: User's clarified question, Extracted academic papers (e.g., ArXiv), Web search content, Uploaded document extracts (if present).

Use all sources collectively and cohesively to produce a single, well-organized report.
            """),
            ("human", "Sources to summarize:\n\n{content}")
        ])
        
        chain = prompt_template | primary_llm
        prompt_data = {"content": content}
        
        response_content = await safe_llm_invoke(chain, prompt_data)
        
        if response_content and not response_content.startswith("Error:"):
            state["source_summaries"] = [response_content]
            print("✍️ Sources summarized successfully")
        else:
            # Fallback: use raw sources
            state["source_summaries"] = all_sources[:3]  # Limit to first 3 sources
            state["errors"].append("Source summarization had issues, using raw sources")
            print("⚠️ Source summarization failed, using raw sources")
            
    except Exception as e:
        print(f"❌ Source summarization error: {e}")
        # Fallback to raw sources
        all_sources = []
        if state.get("uploaded_documents_raw"):
            all_sources.extend(state["uploaded_documents_raw"])
        if state.get("internet_results_raw"):
            all_sources.extend(state["internet_results_raw"])
        if state.get("arxiv_results_raw"):
            all_sources.extend(state["arxiv_results_raw"])
        
        state["source_summaries"] = all_sources[:3] if all_sources else ["Error processing sources"]
        state["errors"].append(f"Source summarization failed: {str(e)}")
    
    return state


async def report_generation_node(state: ResearchState) -> ResearchState:
    """Generate comprehensive research report with detailed content"""
    print("📝 Step 6: Generating Comprehensive Research Report...")
    state["processing_stage"] = "report_generation"
    
    try:
        question = state.get("formatted_question", state["original_question"])
        summaries = state.get("source_summaries", [])
        
        if not summaries:
            state["research_summary"] = f"""
            <div style="max-width: 1000px; margin: 0 auto; font-family: Arial, sans-serif; line-height: 1.6;">
                <h1 style="color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px;">Research Report: {question}</h1>
                <div style="background: #f8f9fa; padding: 20px; border-left: 4px solid #e74c3c; margin: 20px 0;">
                    <h3 style="color: #e74c3c; margin-top: 0;">No Content Available</h3>
                    <p>No detailed content could be retrieved for analysis on this topic.</p>
                </div>
            </div>
            """
            return state
        
        # Combine all content
        combined_content = "\n<hr/>\n".join(summaries)
        
        # Limit content for LLM processing
        if len(combined_content) > 18000:  # Larger limit for detailed reports
            combined_content = combined_content[:18000] + "\n\n[Content truncated for processing...]"
        
        prompt_template = ChatPromptTemplate.from_messages([
    ("system", """You are an expert research analyst generating comprehensive, detailed reports.

        CRITICAL INSTRUCTIONS:
        1. Create a comprehensive research report with substantial detail
        2. Use ALL provided content to write thorough analysis
        3. Structure with clear sections: Summary, Methodology Overview, Elaborated System Architecture, Key findings, Real world example in layman terms, Concise conclusion
        4. Include specific details from the research papers and sources
        5. Return ONLY HTML content - NO markdown code blocks, NO ```html tags, NO DOCTYPE/html/head/body tags
        6. Start directly with the styled container div
        7. Use dark theme colors matching the interface
        8. DO NOT include URLs, publication details, or citation metadata

        HTML STRUCTURE - Return exactly this format:
        <div style="background: #2a2d3a; color: #e4e6ea; max-width: 1000px; margin: 0 auto; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; padding: 20px; border-radius: 8px;">
        <h1 style="color: #60a5fa; border-bottom: 2px solid #60a5fa; padding-bottom: 10px; margin-bottom: 20px;">Title</h1>
        <h2 style="color: #34d399; margin-top: 30px; margin-bottom: 15px; border-left: 4px solid #34d399; padding-left: 15px;">Section Title</h2>
        <p style="margin-bottom: 15px; text-align: justify;">Content here...</p>
        <h3 style="color: #fbbf24; margin-top: 20px; margin-bottom: 10px;">Subsection</h3>
        <ul style="margin-left: 20px; margin-bottom: 15px;">
            <li style="margin-bottom: 5px;">List item</li>
        </ul>
        </div>

        DARK THEME COLORS TO USE:
        - Background: #2a2d3a
        - Text: #e4e6ea  
        - Headings h1: #60a5fa (blue)
        - Headings h2: #34d399 (green)
        - Headings h3: #fbbf24 (yellow)
        - Borders/accents: #4b5563
        - Cards/sections: #374151

        Return ONLY the HTML div content, no other text or formatting."""),
            
            ("human", "Research Question: {question}\n\nDetailed Source Content:\n{content}\n\nGenerate a comprehensive, detailed research report using ALL the provided content:")
        ])
        
        chain = prompt_template | primary_llm
        prompt_data = {"question": question, "content": combined_content}
        
        response_content = await safe_llm_invoke(chain, prompt_data)
        print("generated report=====", response_content)
        if response_content and len(response_content) > 100:
            state["research_summary"] = response_content
            print("📝 Comprehensive research report generated successfully")
        else:
            # Create detailed fallback report with proper HTML structure
            fallback_sections = []
            fallback_sections.append(f"""
            <div style="max-width: 1000px; margin: 0 auto; font-family: Arial, sans-serif; line-height: 1.6;">
                <h1 style="color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px;">Research Report: {question}</h1>
                
                <div style="background: #e8f4f8; padding: 20px; border-left: 4px solid #3498db; margin: 20px 0;">
                    <h2 style="color: #2c3e50; margin-top: 0;">Executive Summary</h2>
                    <p>This comprehensive report analyzes available research on: <strong>{question}</strong></p>
                    <p>The analysis draws from multiple sources including academic papers, web research, and uploaded documents to provide a thorough understanding of the topic.</p>
                </div>
                
                <div style="background: #f8f9fa; padding: 20px; margin: 20px 0; border-radius: 5px;">
                    <h2 style="color: #2c3e50;">Detailed Findings</h2>
            """)
            
            # Include actual content from sources
            for i, content in enumerate(summaries[:3]):
                fallback_sections.append(f"""
                    <div style="margin: 15px 0; padding: 15px; background: white; border-left: 3px solid #3498db;">
                        <h3 style="color: #34495e;">Research Source {i+1}</h3>
                        <div>{content}</div>
                    </div>
                """)
            
            fallback_sections.append(f"""
                </div>
                
                <div style="background: #d5e8d4; padding: 20px; border-left: 4px solid #27ae60; margin: 20px 0;">
                    <h2 style="color: #2c3e50;">Conclusions</h2>
                    <p>Based on the comprehensive analysis of available sources, this research provides valuable insights into {question}.</p>
                    <p>The findings demonstrate the complexity and multifaceted nature of this topic, requiring continued research and analysis.</p>
                </div>
                
                <div style="background: #fdf2e9; padding: 15px; border-left: 4px solid #f39c12; margin: 20px 0;">
                    <h3 style="color: #e67e22; margin-top: 0;">Research Methodology</h3>
                    <p>This report synthesizes information from {len(summaries)} primary sources including academic papers, web research, and document analysis.</p>
                </div>
            </div>
            """)

            state["research_summary"] = "\n".join(fallback_sections)
            print("⚠️ Using detailed structured fallback report")
            
    except Exception as e:
        print(f"❌ Report generation error: {e}")
        # Emergency fallback with proper HTML structure
        state["research_summary"] = f"""
        <div style="max-width: 1000px; margin: 0 auto; font-family: Arial, sans-serif; line-height: 1.6;">
            <h1 style="color: #2c3e50; border-bottom: 3px solid #e74c3c; padding-bottom: 10px;">Research Report: {question}</h1>
            <div style="background: #f8d7da; padding: 20px; border-left: 4px solid #e74c3c; margin: 20px 0;">
                <h3 style="color: #721c24; margin-top: 0;">Report Generation Error</h3>
                <p>An error occurred while generating the comprehensive report: {str(e)}</p>
                <p>Please try again or contact support if the issue persists.</p>
            </div>
        </div>
        """
        state["errors"].append(f"Report generation failed: {str(e)}")
    
    return state


async def qa_specialist_node(state: ResearchState) -> ResearchState:
    """Enhanced QA with full context access and persistent chat history"""
    print("💬 Enhanced Q&A Specialist - full context + persistent chat...")
    state["processing_stage"] = "qa_response"
    
    try:
        user_question = state.get("user_question", "")
        formatted_question = state.get("formatted_question", "")
        if not user_question:
            state["specialist_response"] = "No question provided."
            return state
        follow_up_type = state.get("follow_up_question", "")
        # Gather ALL available context sources
        all_sources = []
        
        # Add uploaded documents

        # Add source summaries
        if state.get("source_summaries"):
            all_sources.extend(state["source_summaries"])
            
        # Add main research report
        if state.get("research_summary"):
            all_sources.append(f"## Main Research Report\n{state['research_summary']}")
        
        # Combine all knowledge sources
        combined_context = "\n\n---SOURCE BREAK---\n\n".join(all_sources)
        
        # Format chat history for context
        chat_history = state.get("chat_history", [])
        chat_context = ""
        
        # Include recent chat history (last 10 exchanges)
        recent_chat = chat_history[-5:] if len(chat_history) > 5 else chat_history
        for msg in recent_chat:
            role = msg.get("role", "user")
            content = msg.get("content", "")
            prefix = "👤 User:" if role == "user" else "🤖 Assistant:"
            chat_context += f"{prefix} {content}\n\n"
        
        # Limit context size for LLM
        if len(combined_context) > 7000:
            combined_context = combined_context[:7000] + "\n\n[Knowledge base truncated...]"
        print(f"followup_type=========={follow_up_type}")
        if follow_up_type == "research_followup":
            prompt_data = {
            "chat_history": chat_context,
            "context": combined_context,
            "question": user_question,
            "formatted_question": formatted_question
        }
            qa_prompt = """
You are an expert research assistant with access to uploaded documents, internet search, ArXiv papers, previous chat history, and user-provided context.

INSTRUCTIONS:
1. Use the full conversation history to maintain continuity and context.
2. Reference only available, verifiable sources from the accessible knowledge base. If detailed information is explicitly requested — or necessary for clarity — provide thorough, step-by-step explanations with examples where applicable.
3. If data is missing or not found in the provided sources, state this limitation clearly.
4. Format all responses strictly using the HTML structure below. Do NOT use Markdown (e.g., no **bold**, use <strong>). Use semantic HTML: <h1>, <h2>, <h3>, <p>, <ul>, <ol>, <em>, and <br> only where needed.
5. Ask the user if they want to go deeper into the topic by mentioning the specific section or aspect they are interested in.

HTML STRUCTURE - Return exactly this format:
<div style="background: #2a2d3a; color: #e4e6ea; max-width: 1000px; margin: 0 auto; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; padding: 20px; border-radius: 8px;">
  <h1 style="color: #60a5fa; border-bottom: 2px solid #60a5fa; padding-bottom: 10px; margin-bottom: 20px;">Title</h1>
  <h2 style="color: #34d399; margin-top: 30px; margin-bottom: 15px; border-left: 4px solid #34d399; padding-left: 15px;">Section Title</h2>
  <p style="margin-bottom: 15px; text-align: justify;">Content here...</p>
  <h3 style="color: #fbbf24; margin-top: 20px; margin-bottom: 10px;">Subsection</h3>
  <ul style="margin-left: 20px; margin-bottom: 15px;">
    <li style="margin-bottom: 5px;">List item</li>
  </ul>
</div>

DARK THEME COLORS TO USE:
- Background: #2a2d3a
- Text: #e4e6ea  
- Headings h1: #60a5fa (blue)
- Headings h2: #34d399 (green)
- Headings h3: #fbbf24 (yellow)
- Borders/accents: #4b5563
- Cards/sections: #374151

Do not use Markdown syntax like **bold**; instead, use proper HTML tags such as <strong> for bold, <em> for italics, <ul>/<ol> for lists, and <br> only when line breaks are necessary within the same block.

"""

        elif follow_up_type == "chat_followup":
            prompt_data = {
            "chat_history": chat_context,
            "context": combined_context,
            "question": user_question,
            "formatted_question": ""
        }
            qa_prompt = """
You are an expert research assistant with access to uploaded documents, internet search, ArXiv papers, previous chat history, and user-provided context.

INSTRUCTIONS:
1. Answer clearly and concisely using ONLY verified information from accessible sources. If something is missing, explicitly state the limitation.
2. Use full conversation history to maintain context and coherence across follow-ups.
3. Format all responses strictly using valid HTML. Do NOT use Markdown (e.g., no **bold**, use <strong>), and apply proper semantic tags: <h1>, <h2>, <h3>, <p>, <ul>, <ol>, <em>, <br>.
4. Be casual and friendly, but maintain professionalism. Use a tone suitable for an intelligent non-expert audience and maintain a human like conversational style.
5. Respond inside this exact structure:

<div>
  <h1>Title</h1>
  <h2>Section Title</h2>
  <p>Content here...</p>
  <h3>Subsection</h3>
  <ul>
    <li>List item</li>
  </ul>
</div>

6. Be precise, well-reasoned, and complete. If detailed explanation is asked, break down the logic step-by-step with examples.

Always conclude with a helpful follow-up suggestion.
"""

        prompt_template = ChatPromptTemplate.from_messages([
            ("system", qa_prompt),
            ("human", """CONVERSATION HISTORY:
{chat_history}

COMPREHENSIVE KNOWLEDGE BASE:
{context}

MY ACTUAL QUESTION: {question}
             
AI FORMATTED QUESTION: {formatted_question}
(Only use this if this is related to the user's question and the chat context, otherwise ignore it)
Please provide a detailed answer using the conversation history and available knowledge sources:""")
        ])
        
        chain = prompt_template | primary_llm
        
        print(f"-=-=-=-=-=-=Prompt Data-=-=-=-=-=: {prompt_data}")
        response_content = await safe_llm_invoke(chain, prompt_data)
        
        if response_content and not response_content.startswith("Error:"):
            state["specialist_response"] = response_content
            
            # Update chat history
            chat_history.append({"role": "user", "content": user_question})
            chat_history.append({"role": "assistant", "content": response_content})
            state["chat_history"] = chat_history
            
            # Update session metadata
            metadata = state.get("session_metadata", {})
            metadata["last_updated"] = datetime.now().isoformat()
            metadata["message_count"] = len(chat_history)
            state["session_metadata"] = metadata
            
            print(f"💬 Enhanced QA response generated. Chat history: {len(chat_history)} messages.")
        else:
            fallback_response = "I apologize, but I'm having trouble processing your question right now. Please try rephrasing or ask a different question."
            state["specialist_response"] = fallback_response
            state["errors"].append("Enhanced QA processing had issues")
            
    except Exception as e:
        print(f"❌ Enhanced QA specialist error: {e}")
        state["specialist_response"] = f"An error occurred while processing your question: {str(e)}"
        state["errors"].append(f"Enhanced QA specialist failed: {str(e)}")
    
    return state


# === Graph Routers (Conditional Logic) ===
def route_entry(state: ResearchState) -> Literal["document_processor", "formatter"]:
    """Route based on whether documents are uploaded"""
    print("🚦 Entry Router: Checking for documents...")
    try:
        research_summary  = state.get("research_summary", "").strip()
        if state.get("uploaded_documents_raw") and len(research_summary) == 0:
            print("🚦 -> Documents found, proceeding to document processing")
            return "document_processor"
        else:
            print("🚦 -> No documents, proceeding directly to formatter")
            return "formatter"
    except Exception as e:
        print(f"❌ Entry routing error: {e}")
        return "formatter"  # Safe default


def evaluate_search_results(state: ResearchState) -> Literal["summarize_sources", "handle_insufficient_data"]:
    """Evaluate if we have sufficient search results"""
    print("🚦 Evaluating Search Results...")
    try:
        has_internet = bool(state.get("internet_results_raw"))
        has_arxiv = bool(state.get("arxiv_results_raw"))
        has_docs = bool(state.get("uploaded_documents_raw"))
        
        # If we have any substantial results, proceed to summarization
        if has_internet or has_arxiv or has_docs:
            print("🚦 -> Sufficient data found, proceeding to summarization")
            return "summarize_sources"
        else:
            print("🚦 -> Insufficient data, handling gracefully")
            return "handle_insufficient_data"
    except Exception as e:
        print(f"❌ Search evaluation error: {e}")
        return "handle_insufficient_data"  # Safe default
def formatter_routing_node(state: ResearchState) -> Literal["formatter", "enhanced_formatter"]:
    """
    Route to appropriate formatter based on question type:
    - Initial research queries -> basic formatter 
    - Follow-up questions -> enhanced formatter (with chat history context)
    """
    print("🚦 Formatter Router: Determining question type...")
    
    try:
        # Check if this is a follow-up question
        user_question = state.get("user_question", "").strip()
        chat_history = state.get("chat_history", [])
        
        # Determine if it's a follow-up based on multiple factors
        is_followup = bool(user_question) and len(chat_history) > 0
        
        if is_followup:
            print("🚦 -> Follow-up question detected, using enhanced formatter")
            return "enhanced_formatter"
        else:
            print("🚦 -> Initial research query, using basic formatter")
            return "formatter"
            
    except Exception as e:
        print(f"❌ Formatter routing error: {e}")
        return "formatter"  # Safe default
def route_research_or_followup(
    state: ResearchState
) -> Literal["report_generator", "enhanced_qa"]:
    """
    Decide whether to:
      • run full research (report_generator)  –– initial query
      • answer via enhanced QA               –– follow-up query

    Heuristic:
    1. If a research_summary already exists AND a user_question is present
       → treat as follow-up → enhanced_qa.
    2. Otherwise → treat as fresh research → report_generator.
    """
    print("🚦 Research/Followup Router: Determining query type...")

    try:
        user_question     = state.get("user_question", "").strip()
        research_summary  = state.get("research_summary", "").strip()
        print("User Question:", user_question)
        print('--------RESEARCH SUMMARYY TYPE---------',type(research_summary))
        is_followup = bool(user_question) and bool(research_summary)

        if is_followup:
            print("🚦 -> Follow-up detected (summary exists), routing to enhanced QA")
            return "qa_specialist"
        else:
            print("🚦 -> No prior summary, routing to report generator")
            return "report_generator"

    except Exception as e:
        print(f"❌ Research/followup routing error: {e}")
        return "report_generator"


# === Graph Construction ===
def create_research_workflow():
    """Create and compile the research workflow"""
    try:
        workflow = StateGraph(ResearchState)
        
        # Add nodes
        workflow.add_node("document_processor", document_processor_node)
        workflow.add_node("query_generator", query_generator_node)
        workflow.add_node("formatter", formatter_node)
        workflow.add_node("internet_search", internet_search_node)
        workflow.add_node("arxiv_search", arxiv_search_node)
        workflow.add_node("summarize_sources", summarize_sources_node)
        workflow.add_node("handle_insufficient_data", handle_insufficient_data_node)
        workflow.add_node("report_generator", report_generation_node)
        workflow.add_node("qa_specialist", qa_specialist_node)
        workflow.add_node("enhanced_formatter", enhanced_formatter_node)
        # Conditional Entry Point
        workflow.add_conditional_edges(
            START,
            route_entry,
            {
                "document_processor": "document_processor",
                "formatter": "formatter"
            }
        )
        
        # Document-First Path
        workflow.add_edge("document_processor", "query_generator")
        workflow.add_conditional_edges(
            "query_generator",
            formatter_routing_node,
            {
                "enhanced_formatter": "enhanced_formatter",
                "formatter": "formatter"
            }
        )
        workflow.add_edge("formatter", "internet_search")

        # Main Research Path
        workflow.add_edge('enhanced_formatter', 'internet_search')        
        workflow.add_edge("internet_search", "arxiv_search")
        
        # Conditional Path based on Search Results
        workflow.add_conditional_edges(
            "arxiv_search",
            evaluate_search_results,
            {
                "summarize_sources": "summarize_sources",
                "handle_insufficient_data": "handle_insufficient_data"
            }
        )
        # Connect to Report Generation
        workflow.add_conditional_edges(
            "handle_insufficient_data",
            route_research_or_followup,
            {
                "report_generator": "report_generator",
                "qa_specialist": "qa_specialist"
            }
        )
        workflow.add_conditional_edges(
            "summarize_sources",
            route_research_or_followup,
            {
                "report_generator": "report_generator",
                "qa_specialist": "qa_specialist"
            }
        )
        # Final Nodes
        workflow.add_edge("qa_specialist", END)
        workflow.add_edge("report_generator", END)
        
        research_app = workflow.compile()
        print("✅ Research workflow compiled successfully")
        return research_app
        
    except Exception as e:
        print(f"❌ Workflow creation failed: {e}")
        raise

# === LangGraph Native Visualization ===
def visualize_langgraph_workflows():
    """Use LangGraph's built-in visualization"""
    try:
        import matplotlib.pyplot as plt
        from IPython.display import Image
        
        # Create both workflows
        research_workflow = create_research_workflow()
        
        workflows = [
            ("Research Workflow", research_workflow),
        ]
        
        fig, axes = plt.subplots(1, len(workflows), figsize=(20, 10))
        if len(workflows) == 1:
            axes = [axes]
        
        for i, (title, workflow) in enumerate(workflows):
            try:
                # Use LangGraph's get_graph() method for visualization
                graph = workflow.get_graph()
                
                # Try different visualization methods
                if hasattr(graph, 'draw_mermaid_png'):
                    # Method 1: Mermaid PNG (best option)
                    png_data = graph.draw_mermaid_png()
                    from PIL import Image as PILImage
                    import io
                    img = PILImage.open(io.BytesIO(png_data))
                    axes[i].imshow(img)
                    axes[i].axis('off')
                elif hasattr(graph, 'draw_ascii'):
                    # Method 2: ASCII representation
                    ascii_repr = graph.draw_ascii()
                    axes[i].text(0.05, 0.95, ascii_repr, transform=axes[i].transAxes,
                               fontfamily='monospace', fontsize=8, va='top', ha='left')
                    axes[i].axis('off')
                else:
                    # Fallback: Show available methods
                    methods = [method for method in dir(graph) if 'draw' in method.lower()]
                    axes[i].text(0.5, 0.5, f"{title}\nAvailable methods:\n" + "\n".join(methods), 
                               ha='center', va='center', transform=axes[i].transAxes)
                    axes[i].axis('off')
                
            except Exception as e:
                # Show error info
                axes[i].text(0.5, 0.5, f"{title}\nVisualization Error:\n{str(e)}", 
                           ha='center', va='center', transform=axes[i].transAxes)
                axes[i].axis('off')
            
            axes[i].set_title(title, fontsize=14, fontweight='bold')
        
        plt.tight_layout()
        plt.savefig('langgraph_native_visualization.png', dpi=300, bbox_inches='tight')
        plt.show()
        
        print("✅ LangGraph native visualization generated")
        
        # Also try individual graph outputs
        for title, workflow in workflows:
            try:
                graph = workflow.get_graph()
                print(f"\n=== {title} ===")
                if hasattr(graph, 'draw_ascii'):
                    print("ASCII representation:")
                    print(graph.draw_ascii())
                if hasattr(graph, 'draw_mermaid'):
                    print("Mermaid code:")
                    print(graph.draw_mermaid())
            except Exception as e:
                print(f"Error with {title}: {e}")
        
    except Exception as e:
        print(f"⚠️ LangGraph native visualization failed: {e}")
        print("Available graph methods:")
        try:
            research_workflow = create_research_workflow()
            graph = research_workflow.get_graph()
            methods = [method for method in dir(graph) if not method.startswith('_')]
            print(methods)
        except Exception as inner_e:
            print(f"Could not inspect graph methods: {inner_e}")

# === FastAPI Application Setup ===
app = FastAPI(
    title="DocuMentor API", 
    version="2.0.0",
    description="AI-Powered Research Assistant with Multi-Agent Processing"
)


# Add CORS middleware for development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")


# === Pydantic Models ===
class QueryRequest(BaseModel):
    question: str
    session_id: str


class FollowupRequest(BaseModel):
    question: str
    session_id: str


class NewSessionRequest(BaseModel):
    session_id: str


# === API Endpoints ===
@app.get("/", include_in_schema=False)
async def serve_index():
    """Serve the main HTML page"""
    try:
        return FileResponse("static/index.html")
    except Exception as e:
        print(f"❌ Error serving index: {e}")
        raise HTTPException(status_code=500, detail="Unable to serve main page")


@app.post("/query")
async def process_query(
    request: Request,
    question: str = Form(None),
    session_id: str = Form(None),
    files: List[UploadFile] = File(None)
):
    """Main endpoint for processing research queries"""
    try:
        # Parse request data
        if not question:
            try:
                body = await request.body()
                if body:
                    import json
                    data = json.loads(body.decode())
                    question = data.get("question", "")
                    session_id = data.get("session_id", "")
            except Exception:
                pass
        
        if not question:
            raise HTTPException(status_code=400, detail="Question is required")
        
        if not session_id:
            session_id = f"session_{uuid.uuid4().hex[:8]}"
        
        print(f"🔬 Processing query: {question[:50]}...")
        
        # Process file uploads
        uploaded_documents = []
        processing_errors = []

        if files:
            for file in files:
                if file.filename:
                    try:
                        # Validate file size
                        if file.size and file.size > MAX_DOCUMENT_SIZE:
                            error_msg = f"File '{file.filename}' exceeds {MAX_DOCUMENT_SIZE//1024//1024}MB limit"
                            processing_errors.append(error_msg)
                            continue
                        
                        # Validate file type
                        file_ext = os.path.splitext(file.filename)[1].lower()
                        if file_ext not in ALLOWED_FILE_TYPES:
                            error_msg = f"File type '{file_ext}' not supported. Allowed: {', '.join(ALLOWED_FILE_TYPES)}"
                            processing_errors.append(error_msg)
                            continue
                        
                        # Save file temporarily
                        file_id = str(uuid.uuid4())
                        unique_filename = f"{file_id}{file_ext}"
                        file_path = os.path.join(UPLOAD_DIR, unique_filename)
                        
                        # Save file content
                        with open(file_path, "wb") as buffer:
                            content = await file.read()
                            buffer.write(content)
                        
                        print(f"📁 Processing file: {file.filename} ({file_ext})")
                        
                        # Extract text using enhanced function
                        text_content = safe_extract_text_from_file(file_path)
                        
                        if text_content and text_content.strip() and not text_content.startswith("Error"):
                            uploaded_documents.append(f"Document: {file.filename}\n\n{text_content}")
                            print(f"✅ Successfully processed: {file.filename}")
                        else:
                            if text_content.startswith("Error") or text_content.startswith("Legacy DOC"):
                                processing_errors.append(f"{file.filename}: {text_content}")
                            else:
                                processing_errors.append(f"No readable text found in '{file.filename}'")
                        
                        # Clean up temporary file
                        try:
                            os.remove(file_path)
                        except Exception as cleanup_error:
                            print(f"⚠️ Cleanup warning for {file_path}: {cleanup_error}")
                            
                    except Exception as e:
                        error_msg = f"Failed to process '{file.filename}': {str(e)}"
                        processing_errors.append(error_msg)
                        print(f"❌ {error_msg}")
            
        # Create initial state

        # In process_query function, update initial_state creation:
        initial_state: ResearchState = {
            "original_question": question,
            "generated_question": "",
            "formatted_question": "",
            "session_id": session_id,
            "uploaded_documents_raw": uploaded_documents,
            "internet_results_raw": [],
            "arxiv_results_raw": [],
            "source_summaries": [],
            "research_summary": "",
            "user_question": "",
            "specialist_response": "",
            "follow_up_question": "",
            "errors": [],
            "processing_stage": "initialized",
            "chat_history": [],  # NEW: Empty chat history for new session
            "session_metadata": {  # NEW: Session metadata
                "title": question[:50] + "..." if len(question) > 50 else question,
                "created_at": datetime.now().isoformat(),
                "last_updated": datetime.now().isoformat(),
                "message_count": 0
            }
        }

        if processing_errors:
            initial_state["errors"].extend(processing_errors)
            print(f"⚠️ File processing completed with {len(processing_errors)} warnings/errors")
                # Store session
        session_store[session_id] = initial_state
        
        # Process through workflow
        result = await research_app.ainvoke(initial_state)
        
        # Update session store
        session_store[session_id] = result
        
        # Format response
        sources_used = []
        if result.get("uploaded_documents_raw"):
            sources_used.append("Uploaded Documents")
        if result.get("internet_results_raw"):
            sources_used.append("Internet Search")
        if result.get("arxiv_results_raw"):
            sources_used.append("Academic Papers")
        
        return {
            "session_id": result["session_id"],
            "question": result["original_question"],
            "formatted_question": result.get("formatted_question", ""),
            "response": result.get("research_summary", "No summary generated"),
            "processing_stage": result.get("processing_stage", "completed"),
            "sources_used": sources_used,
            "errors": result.get("errors", [])
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Query processing failed: {e}")
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "error": f"Processing failed: {str(e)}",
                "session_id": session_id if 'session_id' in locals() else "unknown",
                "response": f"<p>I encountered an error processing your request: {str(e)}</p>"
            }
        )

@app.get("/sessions")
async def list_sessions():
    """Get list of all available sessions"""
    try:
        sessions = []
        for session_id, state in session_store.items():
            metadata = state.get("session_metadata", {})
            sessions.append({
                "session_id": session_id,
                "title": metadata.get("title", "Untitled Session"),
                "created_at": metadata.get("created_at", ""),
                "last_updated": metadata.get("last_updated", ""),
                "message_count": metadata.get("message_count", 0),
                "original_question": state.get("original_question", "")
            })
        
        # Sort by last_updated (most recent first)
        sessions.sort(key=lambda x: x.get("last_updated", ""), reverse=True)
        return {"sessions": sessions}
        
    except Exception as e:
        print(f"❌ List sessions failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list sessions: {str(e)}")

@app.get("/session/{session_id}/full")
async def get_full_session(session_id: str):
    """Get complete session data including chat history"""
    try:
        if session_id not in session_store:
            raise HTTPException(status_code=404, detail="Session not found")
        
        state = session_store[session_id]
        return {
            "session_id": session_id,
            "original_question": state.get("original_question", ""),
            "research_summary": state.get("research_summary", ""),
            "chat_history": state.get("chat_history", []),
            "session_metadata": state.get("session_metadata", {}),
            "processing_stage": state.get("processing_stage", ""),
            "errors": state.get("errors", []),
            "sources_available": {
                "uploaded_documents": len(state.get("uploaded_documents_raw", [])),
                "internet_results": len(state.get("internet_results_raw", [])),
                "arxiv_results": len(state.get("arxiv_results_raw", []))
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Get full session failed: {e}")
        raise HTTPException(status_code=500, detail=f"Session retrieval failed: {str(e)}")

@app.post("/followup")
async def process_followup(
    request: Request,
    question: str = Form(None),
    session_id: str = Form(None),
    files: List[UploadFile] = File(None)
):
    """Main endpoint for processing research queries"""
    try:
        # Parse request data
        if not question:
            try:
                body = await request.body()
                if body:
                    import json
                    data = json.loads(body.decode())
                    question = data.get("question", "")
                    session_id = data.get("session_id", "")
            except Exception:
                pass
        
        if not question:
            raise HTTPException(status_code=400, detail="Question is required")
        
        if not session_id:
            session_id = f"session_{uuid.uuid4().hex[:8]}"
        
        print(f"🔬 Processing query: {question[:50]}...")
        state = session_store[session_id].copy()
        state["user_question"] = question
        state['follow_up_question'] = "research_followup"
        # Process through workflow
        result = await research_app.ainvoke(state)
        
        session_store[session_id] = result
        print(f"------Specialist response======: {result.get('specialist_response', 'No response generated')}")
        print(state)
        return {
            "session_id": result["session_id"],
            "user_question": question,
            "specialist_response": result.get("specialist_response", "No response generated"),
            "processing_stage": result.get("processing_stage", "completed"),
            "chat_history": result.get("chat_history", [])  # NEW: Return chat history
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Enhanced follow-up processing failed: {e}")
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "error": f"Enhanced follow-up processing failed: {str(e)}",
                "session_id": session_id if 'session_id' in locals() else "unknown"
            }
        )

@app.post("/followup-chat")
async def process_followup_chat(
    request: Request,
    question: str = Form(None),
    session_id: str = Form(None)
):
    """Enhanced follow-up with full context and chat history"""
    try:
        # Parse request data
        if not question or not session_id:
            try:
                body = await request.body()
                if body:
                    import json
                    data = json.loads(body.decode())
                    question = question or data.get("question", "")
                    session_id = session_id or data.get("session_id", "")
            except Exception:
                pass
        
        if not question:
            raise HTTPException(status_code=400, detail="Question is required")
        if not session_id:
            raise HTTPException(status_code=400, detail="Session ID is required")
        
        print(f"💬 Processing enhanced follow-up: {question[:50]}...")
        
        # Get session state
        if session_id not in session_store:
            return JSONResponse(
                status_code=400,
                content={
                    "error": "Session not found. Please start with a main research query first.",
                    "session_id": session_id
                }
            )
        
        state = session_store[session_id].copy()
        state["user_question"] = question
        state['follow_up_question'] = "chat_followup"
        # Process through enhanced Q&A specialist
        result = await qa_specialist_node(state)
        
        # Update session store with new chat history
        session_store[session_id] = result
        
        return {
            "session_id": result["session_id"],
            "question": question,
            "specialist_response": result.get("specialist_response", "No response generated"),
            "processing_stage": result.get("processing_stage", "completed"),
            "chat_history": result.get("chat_history", [])  # NEW: Return chat history
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Enhanced follow-up processing failed: {e}")
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "error": f"Enhanced follow-up processing failed: {str(e)}",
                "session_id": session_id if 'session_id' in locals() else "unknown"
            }
        )


@app.post("/new-session")
async def create_new_session(request: NewSessionRequest):
    """Create a new research session"""
    try:
        new_session_id = f"session_{uuid.uuid4().hex[:8]}"
        
        # Initialize empty session
        initial_state: ResearchState = {
            "original_question": "",
            "generated_question": "",
            "formatted_question": "",
            "session_id": new_session_id,
            "uploaded_documents_raw": [],
            "internet_results_raw": [],
            "arxiv_results_raw": [],
            "source_summaries": [],
            "research_summary": "",
            "user_question": "",
            "specialist_response": "",
            "errors": [],
            "processing_stage": "new_session"
        }
        
        session_store[new_session_id] = initial_state
        
        return {
            "session_id": new_session_id,
            "message": "New session created successfully"
        }
        
    except Exception as e:
        print(f"❌ Session creation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Session creation failed: {str(e)}")


@app.get("/session/{session_id}")
async def get_session(session_id: str):
    """Get session information"""
    try:
        if session_id not in session_store:
            raise HTTPException(status_code=404, detail="Session not found")
        
        state = session_store[session_id]
        
        return {
            "session_id": session_id,
            "original_question": state.get("original_question", ""),
            "research_summary": state.get("research_summary", ""),
            "processing_stage": state.get("processing_stage", ""),
            "errors": state.get("errors", [])
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Session retrieval failed: {e}")
        raise HTTPException(status_code=500, detail=f"Session retrieval failed: {str(e)}")


@app.delete("/session/{session_id}")
async def delete_session(session_id: str):
    """Delete a research session"""
    try:
        if session_id in session_store:
            del session_store[session_id]
            return {"message": f"Session {session_id} deleted successfully"}
        else:
            raise HTTPException(status_code=404, detail="Session not found")
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Session deletion failed: {e}")
        raise HTTPException(status_code=500, detail=f"Session deletion failed: {str(e)}")


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    try:
        # Test LLM connectivity
        llm_status = "healthy" if primary_llm else "unavailable"
        
        return {
            "status": "healthy",
            "service": "DocuMentor API v2.0.0",
            "llm_status": llm_status,
            "active_sessions": len(session_store),
            "upload_directory": UPLOAD_DIR
        }
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "status": "unhealthy",
                "error": str(e)
            }
        )


# === Error Handlers ===
@app.exception_handler(404)
async def not_found_handler(request: Request, exc: HTTPException):
    """Handle 404 errors by serving the main page"""
    return FileResponse("static/index.html")


@app.exception_handler(ValidationError)
async def validation_exception_handler(request: Request, exc: ValidationError):
    """Handle validation errors"""
    print(f"❌ Validation Error: {exc}")
    return JSONResponse(
        status_code=422,
        content={
            "error": "Validation failed",
            "details": exc.errors(),
            "message": "Please check your request format"
        }
    )


# === Startup and Shutdown Events ===
@app.on_event("startup")
async def startup_event():
    """Initialize the application"""
    print("🚀 Starting DocuMentor API...")
    
    # Initialize LLMs
    if not initialize_llms():
        print("❌ Critical: LLM initialization failed. Some features may not work.")
    
    # Create research workflow
    global research_app
    try:
        research_app = create_research_workflow()
        print("✅ Research workflow initialized")
        
        # ADD THIS: Generate workflow visualization
        try:
            visualize_langgraph_workflows()
        except Exception as e:
            print(f"⚠️ Graph visualization failed: {e}")
        
    except Exception as e:
        print(f"❌ Critical: Workflow initialization failed: {e}")
        raise
    
    print("✅ DocuMentor API startup complete")


@app.on_event("shutdown") 
async def shutdown_event():
    """Cleanup on shutdown"""
    print("🛑 Shutting down DocuMentor API...")
    
    # Clean up temporary files
    try:
        if os.path.exists(UPLOAD_DIR):
            for file in os.listdir(UPLOAD_DIR):
                try:
                    os.remove(os.path.join(UPLOAD_DIR, file))
                except:
                    pass
    except Exception as e:
        print(f"⚠️ Cleanup warning: {e}")
    
    print("✅ Shutdown complete")

# ─── MOCK-TEST ROUTER ────────────────────────────────────────────
mock_router = APIRouter()

class AnswerPayload(BaseModel):
    answers: dict

def _get_session(session_id: str) -> ResearchState:
    if session_id not in session_store:
        raise HTTPException(status_code=404, detail="Session not found")
    return session_store[session_id]

@mock_router.post("/mock-test/{session_id}")
async def generate_mock_test(session_id: str):
    state = _get_session(session_id)
    
    # Build knowledge blob
    blobs = []
    for k in ("uploaded_documents_raw", "internet_results_raw", "arxiv_results_raw", "research_summary"):
        val = state.get(k)
        if val:
            blobs.extend(val if isinstance(val, list) else [val])
    
    if not blobs:
        raise HTTPException(
            status_code=400,
            detail="This session has no research data yet. Please complete a research query first."
        )
    
    context = "\n\n".join(blobs)[:12_000]
    print(f"Context========",context)
    # Fixed prompt template
    prompt_template = ChatPromptTemplate.from_messages([
    ("system", """You are an expert instructional designer and HTML specialist. And you are also a quiz generation speacilist (AI) that creates EXACTLY 10 multiple-choice questions from provided research content.

STRICT OUTPUT RULES:
1. Generate EXACTLY 10 questions - no more, no less
2. Each question must have EXACTLY 4 options (A, B, C, D)
3. Questions must be derived ONLY from the provided context and should directly relate to the core technical concepts, theories, or findings presented in the research material.
4. The order of questions should be from easy to medium to hard (basic to advanced)
5. Output ONLY the HTML structure below - no explanations, reasoning, or extra text
6. One option per question must be marked with data-correct="true"

MANDATORY HTML STRUCTURE:
```html
<div class="quiz-content">
  <div class="question-item">
    <h3>Question text here?</h3>
    <div class="options">
      <label><input type="radio" name="q1" value="A" required> Option A text</label>
      <label><input type="radio" name="q1" value="B" required> Option B text</label>
      <label data-correct="true"><input type="radio" name="q1" value="C" required> Option C text</label>
      <label><input type="radio" name="q1" value="D" required> Option D text</label>
    </div>
  </div>
  <!-- Repeat for q2 through q10 -->
</div>
     
CRITICAL REQUIREMENTS:
- Use name="q1" through name="q10" for each question group
- Include required attribute on ALL radio inputs
- Place data-correct="true" on exactly ONE label per question
- Input element must come BEFORE the span text
- NO ```html or ``` markdown blocks
- NO additional explanatory text
- Return ONLY the HTML div structure above
You MUST strictly replicate this structure for all 10 questions. And should only generate 10 questions with no additional text or formatting and any explanation or any html tags.
"""),
    ("human", "Research Material:\n{context}\n\nGenerate quiz HTML:")
])

    
    prompt_data = {"context": context}
    html = await safe_llm_invoke(prompt_template | primary_llm, prompt_data)
    
    # Store in session
    state["latest_quiz_html"] = html
    state["latest_quiz_time"] = datetime.utcnow().isoformat()
    session_store[session_id] = state
    
    return {"mock_test_html": html}

@mock_router.post("/mock-test/{session_id}/submit")
async def submit_mock_test(session_id: str, payload: AnswerPayload):
    print(f"============🔬 Submitting mock test for session {session_id}...======")
    """
    Evaluate a student's quiz submission WITHOUT a predefined answer key.
    The LLM receives:
      • The quiz HTML that was sent to the student.
      • All research material collected during the session.
      • The student's answers.
    It must grade strictly on the basis of that material and return a fully
    formatted HTML report.
    """
    # ── Retrieve session ──────────────────────────────────────────────
    state = _get_session(session_id)
    answers = payload.answers
    quiz_html = state.get("latest_quiz_html")

    if not quiz_html:
        raise HTTPException(status_code=400,
                            detail="No quiz exists for this session.")

    # ── Build evidence packet ─────────────────────────────────────────
    sources: list[str] = []
    for key in ("uploaded_documents_raw", "internet_results_raw",
                "arxiv_results_raw", "research_summary"):
        val = state.get(key)
        if val:
            sources.extend(val if isinstance(val, list) else [val])

    if not sources:
        raise HTTPException(status_code=400,
                            detail="No background material available; "
                                   "cannot evaluate answers.")

    #evidence_blob: str = "\n\n---SOURCE BREAK---\n\n".join(sources)[:8_000]
    evidence_blob = state.get("source_summaries")
    # ── Normalise student answers to deterministic order ─────────────
    # payload.answers is a dict keyed by question indices ("0", "1", …)
    # Handle both "q1" format and "1" format
    answer_keys = []
    for key in payload.answers.keys():
        if key.startswith('q'):
            answer_keys.append(int(key[1:]))  # Extract number from "q1" -> 1
        else:
            answer_keys.append(int(key))      # Direct number
    ordered_keys = sorted(answer_keys)
    # Fix the answer processing logic
    student_answers_list = []
    for key, value in payload.answers.items():
        if key.startswith('q'):
            q_num = key[1:]  # Remove 'q' prefix
            student_answers_list.append(f"Q{q_num}: {value}")
        else:
            student_answers_list.append(f"Q{key}: {value}")

    student_answers: str = "\n".join(sorted(student_answers_list))



    # ── Prompt the model ──────────────────────────────────────────────
    prompt = ChatPromptTemplate.from_messages([
    ("system","""
     You are an expert academic evaluator. Grade the student's quiz answers and provide a comprehensive performance report.
     Refer the stundent as 'You' and the student answers as 'Your answers'.
EVALUATION PROCESS:
1. Compare each student answer with the correct answer from quiz HTML
2. Use ONLY the provided evidence to justify correctness
3. Award points: Correct = 1 point, Wrong = 0 points
4. Calculate percentage score (total points / 10 × 100)
5. Analyze performance patterns and provide actionable feedback
6. Be casual and friendly, but maintain professionalism. Use a tone suitable for an intelligent non-expert audience and maintain a human like conversational style.

OUTPUT REQUIREMENTS - Return ONLY in this structure:
    Header: Grading Report
        - Include total score and percentage
    Question Breakdown:
        - For each question, clearly state whether the student's answer is Correct or Incorrect.
            * If Correct – provide a concise justification referencing the supplied evidence.
            * If Incorrect – restate the question and the correct answer, explain the reasoning behind it using only the given evidence, and clarify why the student's answer was incorrect.
    Performance Analysis:
        - Summarize the strengths
        - Highlight areas for improvement
    Next Steps:
        - Recommend specific actions to improve understanding
        - Suggest additional resources or study materials
    AI Study Assistant follow-up prompt:
        - The prompt must focus to help the student master their weakness using detailed explanations, and practical layman examples.
        - Mention the topics and concepts the student struggled with based on the quiz results within the prompt.
    ALL OUTPUT MUST BE IN BELOW FORMAT
    HTML STRUCTURE - Return exactly this format:
        <div style="background: #2a2d3a; color: #e4e6ea; max-width: 1000px; margin: 0 auto; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; padding: 20px; border-radius: 8px;">
        <h1 style="color: #60a5fa; border-bottom: 2px solid #60a5fa; padding-bottom: 10px; margin-bottom: 20px;">Title</h1>
        <h2 style="color: #34d399; margin-top: 30px; margin-bottom: 15px; border-left: 4px solid #34d399; padding-left: 15px;">Section Title</h2>
        <p style="margin-bottom: 15px; text-align: justify;">Content here...</p>
        <h3 style="color: #fbbf24; margin-top: 20px; margin-bottom: 10px;">Subsection</h3>
        <ul style="margin-left: 20px; margin-bottom: 15px;">
            <li style="margin-bottom: 5px;">List item</li>
        </ul>
        </div>

        DARK THEME COLORS TO USE:
        - Background: #2a2d3a
        - Text: #e4e6ea  
        - Headings h1: #60a5fa (blue)
        - Headings h2: #34d399 (green)
        - Headings h3: #fbbf24 (yellow)
        - Borders/accents: #4b5563
        - Cards/sections: #374151
     
    RETURN NOTHING ELSE
     """
),

    ("human",
     f"QUIZ HTML:\n{quiz_html}\n\n"
     f"EVIDENCE (Correct answers, supporting material):\n{evidence_blob}\n\n"
     f"STUDENT ANSWERS:\n{student_answers}\n\n"
     "Grade the quiz submission step-by-step and produce a complete grading report in the specified HTML format. Ensure to follow the evaluation process and output requirements strictly.")
])

    print("quiz_html", quiz_html)
    print("evidence_blob", evidence_blob)
    print("student_answers", student_answers)
    print("=====🔬 Generating mock test evaluation report...======")
    chain = prompt | primary_llm
    evaluation_html = await safe_llm_invoke(chain, {"quiz_html": quiz_html, "student_answers": student_answers, "evidence_blob": evidence_blob})
    print("============ Mock test evaluation report generated! HEMANTHHH======", evaluation_html)
    # ── Persist evaluation for audit/history (optional) ───────────────
    # ---- 2. Add to chat_history -----------------------------------------
    # chat_entry = {
    #     "role": "assistant",
    #     "content": evaluation_html,
    #     "type": "quiz_evaluation",
    #     "timestamp": datetime.utcnow().isoformat(timespec="seconds") + "Z"
    # }
    # state.setdefault("chat_history", []).append(chat_entry)
    # state.setdefault("quiz_history", []).append(
    #     {"submitted": datetime.utcnow().isoformat(timespec="seconds") + "Z",
    #      "answers": evaluation_html,
    #      "evaluation": evaluation_html}
    # )

    # ---- 3. Return in follow-up-style envelope --------------------------
    return {
        "session_id": session_id,
        "question": "📊 Mock-test submitted",
        "specialist_response": evaluation_html,
        "processing_stage": "quiz_evaluated",
        "chat_history": state["chat_history"]
    }


app.include_router(mock_router)  # Remove the /api prefix
# === Development Server ===
if __name__ == "__main__":
    print("🔬 DocuMentor Research Assistant")
    print("=" * 50)
    
    # Initialize before running
    if not initialize_llms():
        print("❌ Cannot start without LLM initialization")
        exit(1)
    
    try:
        research_app = create_research_workflow()
        print("✅ All systems initialized")
        
        # Run the server
        uvicorn.run(
            app, 
            host="127.0.0.1", 
            port=8000,
            log_level="info",
            reload=False  # Set to True for development
        )
        
    except Exception as e:
        print(f"❌ Startup failed: {e}")
        exit(1)